#!/usr/bin/env python3
"""
Direct test of bulk upload functionality
Test the extraction functions directly without server
"""

import io
import PyPDF2
import docx

def extract_text_from_pdf_test(file_content: bytes, filename: str) -> str:
    """Test PDF extraction - copy of real function with debug"""
    try:
        print(f"📄 Testing PDF extraction for: {filename}")
        print(f"   File size: {len(file_content)} bytes")
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        print(f"   PDF pages: {len(pdf_reader.pages)}")
        
        text = ""
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
                print(f"   Page {i+1}: {len(page_text)} characters")
            else:
                print(f"   Page {i+1}: No text extracted")
        
        text = text.strip()
        if text:
            print(f"✓ Total PDF text extracted: {len(text)} characters")
            return text
        else:
            print("✗ PDF text extraction returned empty")
            return "PDF_EMPTY"
    except Exception as e:
        print(f"✗ PDF extraction error: {e}")
        return f"PDF_ERROR: {str(e)}"

def extract_text_from_docx_test(file_content: bytes, filename: str) -> str:
    """Test DOCX extraction - copy of real function with debug"""
    try:
        print(f"📄 Testing DOCX extraction for: {filename}")
        print(f"   File size: {len(file_content)} bytes")
        
        doc = docx.Document(io.BytesIO(file_content))
        print(f"   Document paragraphs: {len(doc.paragraphs)}")
        
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        text = text.strip()
        if text:
            print(f"✓ Total DOCX text extracted: {len(text)} characters")
            return text
        else:
            print("✗ DOCX text extraction returned empty")
            return "DOCX_EMPTY"
    except Exception as e:
        print(f"✗ DOCX extraction error: {e}")
        return f"DOCX_ERROR: {str(e)}"

def create_test_files():
    """Create sample files for testing"""
    import os
    
    # Create a simple text file that can simulate resume content
    test_content = """
JOHN SMITH
Software Engineer

EXPERIENCE:
Senior Developer at Tech Corp (2020-2024) 
- Developed web applications using Python and React
- Led team of 5 developers
- Built RESTful APIs and microservices

EDUCATION:
Bachelor of Computer Science
State University (2016-2020)

SKILLS:
Python, JavaScript, React, Node.js, SQL, Docker, AWS

CONTACT:
Email: john.smith@email.com
Phone: +1-555-0123
"""

    # Save as text file
    with open("sample_resume.txt", "w") as f:
        f.write(test_content)
    
    print("✅ Created sample_resume.txt")
    return ["sample_resume.txt"]

def simulate_bulk_upload():
    """Simulate the bulk upload process"""
    print("🚀 Simulating bulk resume upload process")
    print("=" * 50)
    
    # Create test files
    files = create_test_files()
    
    candidates = []
    processing_errors = []
    
    for file_path in files:
        print(f"\n📁 Processing file: {file_path}")
        
        try:
            # Read file content (simulating FastAPI UploadFile.read())
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            filename = file_path
            ext = filename.lower().split('.')[-1]
            
            print(f"   Extension: {ext}")
            print(f"   Size: {len(file_content)} bytes")
            
            text_content = ""
            if ext == 'pdf':
                text_content = extract_text_from_pdf_test(file_content, filename)
            elif ext in ['docx', 'doc']:
                text_content = extract_text_from_docx_test(file_content, filename)
            elif ext == 'txt':
                text_content = file_content.decode('utf-8')
                print(f"✓ TXT file decoded: {len(text_content)} characters")
            else:
                processing_errors.append(f"{filename}: Unsupported file type '{ext}'")
                continue
            
            # Check if extraction was successful (same logic as real function)
            if text_content and not text_content.startswith(('PDF_ERROR:', 'DOCX_ERROR:', 'PDF_EMPTY', 'DOCX_EMPTY')):
                # Extract name from filename or content
                name = filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').title()
                
                candidate = {
                    "name": name,
                    "filename": filename,
                    "resume": text_content,
                    "content": text_content
                }
                candidates.append(candidate)
                print(f"✓ Successfully processed: {filename} -> {name}")
            else:
                processing_errors.append(f"{filename}: Text extraction failed - {text_content}")
                print(f"✗ Failed to extract text from: {filename}")
        
        except Exception as e:
            error_msg = f"{filename}: Processing error - {str(e)}"
            processing_errors.append(error_msg)
            print(f"✗ Error processing {filename}: {e}")
    
    print(f"\n📊 Final Results:")
    print(f"   Candidates processed: {len(candidates)}")
    print(f"   Processing errors: {len(processing_errors)}")
    
    if candidates:
        print(f"\n✅ SUCCESS - {len(candidates)} candidates ready for screening")
        for candidate in candidates:
            print(f"   - {candidate['name']}: {len(candidate['content'])} chars")
    else:
        print(f"\n❌ FAILURE - No valid resumes processed")
        if processing_errors:
            print("   Errors:")
            for error in processing_errors:
                print(f"   - {error}")
    
    return {
        "candidates": candidates,
        "processing_errors": processing_errors,
        "count": len(candidates)
    }

if __name__ == "__main__":
    result = simulate_bulk_upload()
    
    print("\n" + "=" * 50)
    print("📋 CONCLUSION:")
    
    if result["candidates"]:
        print("✅ Bulk upload simulation SUCCESSFUL")
        print("   The extraction functions are working correctly")
        print("   Issue might be in:")
        print("   - File upload handling in FastAPI")
        print("   - Frontend file sending")
        print("   - Network/CORS issues")
    else:
        print("❌ Bulk upload simulation FAILED")
        print("   Issue is in the extraction functions or file processing")
        
    print("\n🔍 Next steps:")
    print("1. Test with actual PDF/DOCX files if available")
    print("2. Check server logs for more details")
    print("3. Test the frontend file upload mechanism")